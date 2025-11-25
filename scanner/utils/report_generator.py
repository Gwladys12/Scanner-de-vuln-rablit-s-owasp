import os
import json
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

def generate_report(run_id, report_format="pdf"):
    base_dir = os.path.join("results", run_id)
    input_file = os.path.join(base_dir, "report.json")

    if not os.path.exists(input_file):
        raise FileNotFoundError("Rapport JSON introuvable")

    with open(input_file, encoding="utf-8") as f:
        data = json.load(f)

    # ✅ Correction : gérer si data est une liste ou un dictionnaire
    if isinstance(data, list):
        vulnerabilities = data
    elif isinstance(data, dict) and "vulnerabilities" in data:
        vulnerabilities = data["vulnerabilities"]
    else:
        vulnerabilities = []

    output_dir = os.path.join("scanner", "reports", report_format)
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    output_file = os.path.join(output_dir, f"report_{run_id}_{timestamp}.{report_format}")

    if report_format == "pdf":
        c = canvas.Canvas(output_file, pagesize=A4)
        c.drawString(50, 800, f"Rapport d'analyse OWASP ZAP - {run_id}")
        c.drawString(50, 780, f"Date : {timestamp}")
        c.drawString(50, 760, f"Nombre de vulnérabilités : {len(vulnerabilities)}")
        c.drawString(50, 740, "-"*50)
        y = 720
        for v in vulnerabilities:
            c.drawString(50, y, f"{v.get('type', 'Inconnue')} ({v.get('severity', 'N/A')}): {v.get('description', '')[:70]}")
            y -= 20
            if y < 100:
                c.showPage()
                y = 780
        c.save()

    elif report_format == "docx":
        doc = Document()
        doc.add_heading(f"Rapport OWASP ZAP - {run_id}", 0)
        doc.add_paragraph(f"Date : {timestamp}")
        doc.add_paragraph(f"Nombre de vulnérabilités : {len(vulnerabilities)}")
        for v in vulnerabilities:
            doc.add_heading(v.get("type", "Inconnue"), level=2)
            doc.add_paragraph(f"Criticité : {v.get('severity', 'N/A')}")
            doc.add_paragraph(v.get("description", ""))
        doc.save(output_file)

    elif report_format == "html":
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("<h1>Rapport OWASP ZAP</h1>")
            f.write(f"<p>ID: {run_id}</p><hr>")
            for v in vulnerabilities:
                f.write(f"<h3>{v.get('type', 'Inconnue')} ({v.get('severity', 'N/A')})</h3><p>{v.get('description', '')}</p><hr>")

    return output_file
